"""Word document scanning and identifier block replacement."""

from __future__ import annotations

from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path
import re
import shutil
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from pmid2endnote.endnote import make_temporary_citation
from pmid2endnote.errors import InputDocumentError, WordProcessingError
from pmid2endnote.models import IdentifierKind, ReferenceRecord
from pmid2endnote.pubmed import PubMedRecord
from pmid2endnote.references import reference_from_pubmed
from pmid2endnote.scanner import (
    IdentifierBlock,
    extract_unique_identifiers,
    extract_unique_pmids,
    scan_text,
)


IdentifierKey = tuple[IdentifierKind, str]

REFERENCE_SECTION_HEADINGS = {
    "references",
    "bibliography",
    "works cited",
    "literature cited",
    "references cited",
}

REFERENCE_HEADING_PREFIX_RE = re.compile(
    r"""
    ^
    (?:
        \d+
        |
        [ivxlcdm]+
    )
    [.)]?
    \s+
    """,
    re.IGNORECASE | re.VERBOSE,
)


@dataclass(frozen=True)
class ReplacementBlock:
    """One Word-text range to replace, potentially containing mixed identifiers."""

    original_text: str
    start: int
    end: int
    blocks: tuple[IdentifierBlock, ...]
    source: str

    @property
    def identifier_keys(self) -> tuple[IdentifierKey, ...]:
        keys: list[IdentifierKey] = []
        for block in self.blocks:
            keys.extend((block.kind, identifier) for identifier in block.identifiers)
        return tuple(keys)

    @property
    def pmids(self) -> tuple[str, ...]:
        return tuple(value for kind, value in self.identifier_keys if kind == "pmid")

    @property
    def dois(self) -> tuple[str, ...]:
        return tuple(value for kind, value in self.identifier_keys if kind == "doi")

    @property
    def kind(self) -> str:
        kinds = {kind for kind, _ in self.identifier_keys}
        return next(iter(kinds)) if len(kinds) == 1 else "mixed"


@dataclass(frozen=True)
class TextLocation:
    """Location of a paragraph-like text container in a Word document."""

    part: str
    paragraph_index: int
    comment_id: int | None = None

    def as_report_dict(self) -> dict[str, Any]:
        location = {
            "part": self.part,
            "paragraph_index": self.paragraph_index,
        }
        if self.comment_id is not None:
            location["comment_id"] = self.comment_id
        return location


@dataclass(frozen=True)
class ScannableParagraph:
    """A paragraph-like Word location with reference-section skip state."""

    paragraph: Paragraph
    location: TextLocation
    skipped_by_reference_section: bool = False


@dataclass(frozen=True)
class ReferenceSectionContext:
    """Reference-section metadata used for body/table/comment filtering."""

    paragraphs: list[ScannableParagraph]
    reference_section_start: dict[str, Any] | None
    anchor_locations: dict[int, tuple[Paragraph, TextLocation]]
    skipped_comment_ids: set[int]


@dataclass(frozen=True)
class DocumentPmidOccurrence:
    """An identifier block found in a Word paragraph-like location."""

    block: IdentifierBlock
    location: TextLocation


@dataclass(frozen=True)
class ReplacementOptions:
    """Options that control Word replacement behavior."""

    include_tables: bool = True
    include_comments: bool = True
    include_headers: bool = False
    include_footers: bool = False
    include_footnotes: bool = False
    dry_run: bool = False
    keep_pmid_text: bool = False
    mark_unresolved: bool = False
    scan_parenthetical_pmids: bool = False
    scan_dois: bool = True
    scan_bare_dois: bool = False
    skip_reference_section: bool = True
    create_backup: bool = False


@dataclass(frozen=True)
class ScanResult:
    """PMID scanning result for a Word document."""

    occurrences: list[DocumentPmidOccurrence]
    unique_pmids: list[str]
    unique_identifiers: list[IdentifierKey]
    warnings: list[str]
    skipped_identifiers: list[dict[str, Any]]
    reference_section_start: dict[str, Any] | None


@dataclass(frozen=True)
class ReplacementResult:
    """Replacement result for a Word document."""

    replacements: list[dict[str, Any]]
    warnings: list[str]
    backup_path: Path | None = None


def default_output_path(input_docx: Path) -> Path:
    return input_docx.with_name(f"{input_docx.stem}.endnote.docx")


def default_nbib_path(input_docx: Path) -> Path:
    return input_docx.with_name(f"{input_docx.stem}.references.nbib")


def default_enw_path(input_docx: Path) -> Path:
    return input_docx.with_name(f"{input_docx.stem}.endnote-import.enw")


def default_report_path(input_docx: Path) -> Path:
    return input_docx.with_name(f"{input_docx.stem}.pmid2endnote.report.json")


def validate_input_docx(path: Path) -> None:
    if not path.exists():
        raise InputDocumentError(f"Input .docx does not exist: {path}")
    if not path.is_file():
        raise InputDocumentError(f"Input path is not a file: {path}")
    if path.suffix.lower() != ".docx":
        raise InputDocumentError(f"Input file must be a .docx document: {path}")


def is_reference_section_heading(text: str) -> bool:
    """Return True for conservative standalone bibliography headings."""

    normalized = " ".join(text.strip().split())
    normalized = normalized.strip(" \t\r\n:;.").lower()
    normalized = REFERENCE_HEADING_PREFIX_RE.sub("", normalized).strip()
    return normalized in REFERENCE_SECTION_HEADINGS


def scan_docx(path: Path, options: ReplacementOptions | None = None) -> ScanResult:
    """Scan a .docx document for PMID blocks."""

    options = options or ReplacementOptions()
    validate_input_docx(path)
    document = _open_document(path)
    warnings: list[str] = []
    occurrences: list[DocumentPmidOccurrence] = []
    skipped_identifiers: list[dict[str, Any]] = []
    context = _reference_section_context(document, options)

    for item in context.paragraphs:
        paragraph = item.paragraph
        location = item.location
        blocks = scan_text(
            paragraph.text,
            scan_parenthetical_pmids=options.scan_parenthetical_pmids,
            scan_dois=options.scan_dois,
            scan_bare_dois=options.scan_bare_dois,
        )
        if not blocks:
            continue
        if item.skipped_by_reference_section:
            skipped_identifiers.extend(
                _skipped_identifier_reports(blocks, location, reason="reference_section")
            )
            continue
        if _paragraph_has_unsafe_fields(paragraph):
            warnings.append(
                "Skipped PMID block in paragraph with field or hidden text content at "
                f"{location.part} paragraph {location.paragraph_index}."
            )
            continue
        occurrences.extend(DocumentPmidOccurrence(block=block, location=location) for block in blocks)

    if options.include_comments:
        for paragraph, location in _iter_comment_paragraphs(document, options):
            blocks = scan_text(
                paragraph.text,
                scan_parenthetical_pmids=options.scan_parenthetical_pmids,
                scan_dois=options.scan_dois,
                scan_bare_dois=options.scan_bare_dois,
            )
            if not blocks:
                continue
            if location.comment_id in context.skipped_comment_ids:
                skipped_identifiers.extend(
                    _skipped_identifier_reports(blocks, location, reason="reference_section")
                )
                continue
            if _paragraph_has_unsafe_fields(paragraph):
                warnings.append(
                    "Skipped PMID block in paragraph with field or hidden text content at "
                    f"{location.part} paragraph {location.paragraph_index}."
                )
                continue
            occurrences.extend(
                DocumentPmidOccurrence(block=block, location=location) for block in blocks
            )

    for paragraph, location in _iter_header_footer_paragraphs(document, options):
        blocks = scan_text(
            paragraph.text,
            scan_parenthetical_pmids=options.scan_parenthetical_pmids,
            scan_dois=options.scan_dois,
            scan_bare_dois=options.scan_bare_dois,
        )
        if not blocks:
            continue
        if _paragraph_has_unsafe_fields(paragraph):
            warnings.append(
                "Skipped PMID block in paragraph with field or hidden text content at "
                f"{location.part} paragraph {location.paragraph_index}."
            )
            continue
        occurrences.extend(DocumentPmidOccurrence(block=block, location=location) for block in blocks)

    if skipped_identifiers:
        warnings.append(
            f"Skipped {len(skipped_identifiers)} identifier block"
            f"{'' if len(skipped_identifiers) == 1 else 's'} in the detected reference section."
        )

    if options.include_footnotes:
        warnings.append(
            "Footnotes were requested, but python-docx does not expose footnotes for safe "
            "replacement in this implementation."
        )

    return ScanResult(
        occurrences=occurrences,
        unique_pmids=extract_unique_pmids(occurrence.block for occurrence in occurrences),
        unique_identifiers=extract_unique_identifiers(
            occurrence.block for occurrence in occurrences
        ),
        warnings=warnings,
        skipped_identifiers=skipped_identifiers,
        reference_section_start=context.reference_section_start,
    )


def replace_pmids_in_docx(
    *,
    input_docx: Path,
    output_docx: Path,
    records_by_identifier: dict[IdentifierKey, ReferenceRecord] | None = None,
    records_by_pmid: dict[str, PubMedRecord] | None = None,
    options: ReplacementOptions | None = None,
) -> ReplacementResult:
    """Replace identifier blocks with EndNote temporary citations."""

    options = options or ReplacementOptions()
    resolved_records = _coerce_records_by_identifier(
        records_by_identifier=records_by_identifier,
        records_by_pmid=records_by_pmid,
    )
    validate_input_docx(input_docx)
    document = _open_document(input_docx)
    warnings: list[str] = []
    replacements: list[dict[str, Any]] = []
    context = _reference_section_context(document, options)

    for item in context.paragraphs:
        if item.skipped_by_reference_section:
            continue
        paragraph_replacements, paragraph_warnings = _replace_paragraph_blocks(
            paragraph=item.paragraph,
            location=item.location,
            records_by_identifier=resolved_records,
            keep_pmid_text=options.keep_pmid_text,
            mark_unresolved=options.mark_unresolved,
            dry_run=options.dry_run,
            scan_parenthetical_pmids=options.scan_parenthetical_pmids,
            scan_dois=options.scan_dois,
            scan_bare_dois=options.scan_bare_dois,
        )
        replacements.extend(paragraph_replacements)
        warnings.extend(paragraph_warnings)

    if options.include_comments:
        comment_replacements, comment_warnings = _insert_comment_pmids_at_anchors(
            document=document,
            options=options,
            records_by_identifier=resolved_records,
            context=context,
        )
        replacements.extend(comment_replacements)
        warnings.extend(comment_warnings)

    for paragraph, location in _iter_header_footer_paragraphs(document, options):
        paragraph_replacements, paragraph_warnings = _replace_paragraph_blocks(
            paragraph=paragraph,
            location=location,
            records_by_identifier=resolved_records,
            keep_pmid_text=options.keep_pmid_text,
            mark_unresolved=options.mark_unresolved,
            dry_run=options.dry_run,
            scan_parenthetical_pmids=options.scan_parenthetical_pmids,
            scan_dois=options.scan_dois,
            scan_bare_dois=options.scan_bare_dois,
        )
        replacements.extend(paragraph_replacements)
        warnings.extend(paragraph_warnings)

    if options.include_footnotes:
        warnings.append(
            "Footnotes were requested, but python-docx does not expose footnotes for safe "
            "replacement in this implementation."
        )

    _validate_replacement_citations(replacements)

    backup_path: Path | None = None
    if not options.dry_run:
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        if options.create_backup:
            backup_path = _create_backup(input_docx)
        document.save(output_docx)

    return ReplacementResult(
        replacements=replacements,
        warnings=warnings,
        backup_path=backup_path,
    )


def _open_document(path: Path) -> DocxDocument:
    try:
        return Document(path)
    except Exception as exc:  # python-docx raises a mix of package errors here.
        raise InputDocumentError(f"Could not open .docx document: {path}") from exc


def _coerce_records_by_identifier(
    *,
    records_by_identifier: dict[IdentifierKey, ReferenceRecord] | None,
    records_by_pmid: dict[str, PubMedRecord] | None,
) -> dict[IdentifierKey, ReferenceRecord]:
    if records_by_identifier is not None:
        return records_by_identifier
    if records_by_pmid is None:
        return {}
    return {
        ("pmid", pmid): reference_from_pubmed(record)
        for pmid, record in records_by_pmid.items()
    }


def _reference_section_context(
    document: DocxDocument,
    options: ReplacementOptions,
) -> ReferenceSectionContext:
    paragraphs: list[ScannableParagraph] = []
    reference_section_start: dict[str, Any] | None = None
    in_reference_section = False
    anchor_locations: dict[int, tuple[Paragraph, TextLocation]] = {}
    skipped_comment_ids: set[int] = set()

    for paragraph, location in _iter_body_table_paragraphs_in_order(document, options):
        if (
            options.skip_reference_section
            and not in_reference_section
            and is_reference_section_heading(paragraph.text)
        ):
            in_reference_section = True
            reference_section_start = location.as_report_dict()
        paragraphs.append(
            ScannableParagraph(
                paragraph=paragraph,
                location=location,
                skipped_by_reference_section=in_reference_section,
            )
        )
        for comment_id in _comment_ids_in_paragraph(paragraph):
            if in_reference_section and options.skip_reference_section:
                skipped_comment_ids.add(comment_id)
            else:
                anchor_locations.setdefault(comment_id, (paragraph, location))

    return ReferenceSectionContext(
        paragraphs=paragraphs,
        reference_section_start=reference_section_start,
        anchor_locations=anchor_locations,
        skipped_comment_ids=skipped_comment_ids,
    )


def _iter_body_table_paragraphs_in_order(
    document: DocxDocument,
    options: ReplacementOptions,
) -> Iterator[tuple[Paragraph, TextLocation]]:
    body_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document), TextLocation(part="body", paragraph_index=body_index)
            body_index += 1
        elif child.tag == qn("w:tbl") and options.include_tables:
            table = Table(child, document)
            for paragraph in _iter_table_paragraphs(table):
                yield paragraph, TextLocation(part="table", paragraph_index=table_index)
                table_index += 1

def _iter_header_footer_paragraphs(
    document: DocxDocument,
    options: ReplacementOptions,
) -> Iterator[tuple[Paragraph, TextLocation]]:
    if options.include_headers:
        header_index = 0
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph, TextLocation(part="header", paragraph_index=header_index)
                header_index += 1

    if options.include_footers:
        footer_index = 0
        for section in document.sections:
            for paragraph in section.footer.paragraphs:
                yield paragraph, TextLocation(part="footer", paragraph_index=footer_index)
                footer_index += 1


def _iter_scannable_paragraphs(
    document: DocxDocument,
    options: ReplacementOptions,
    *,
    include_comments: bool = True,
) -> Iterator[tuple[Paragraph, TextLocation]]:
    context = _reference_section_context(document, options)
    for item in context.paragraphs:
        if not item.skipped_by_reference_section:
            yield item.paragraph, item.location
    if include_comments:
        for paragraph, location in _iter_comment_paragraphs(document, options):
            if location.comment_id not in context.skipped_comment_ids:
                yield paragraph, location
    yield from _iter_header_footer_paragraphs(document, options)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_comment_paragraphs(
    document: DocxDocument,
    options: ReplacementOptions,
) -> Iterator[tuple[Paragraph, TextLocation]]:
    if not options.include_comments:
        return
    comment_index = 0
    for comment in document.comments:
        comment_id = int(comment.comment_id)
        for paragraph in comment.paragraphs:
            yield paragraph, TextLocation(
                part="comment",
                paragraph_index=comment_index,
                comment_id=comment_id,
            )
            comment_index += 1
        if options.include_tables:
            for table in comment.tables:
                for paragraph in _iter_table_paragraphs(table):
                    yield paragraph, TextLocation(
                        part="comment",
                        paragraph_index=comment_index,
                        comment_id=comment_id,
                    )
                    comment_index += 1


def _paragraph_has_unsafe_fields(paragraph: Paragraph) -> bool:
    xml = paragraph._p.xml
    return any(marker in xml for marker in ("w:fldChar", "w:instrText", "w:vanish"))


def _replace_paragraph_blocks(
    *,
    paragraph: Paragraph,
    location: TextLocation,
    records_by_identifier: dict[IdentifierKey, ReferenceRecord],
    keep_pmid_text: bool,
    mark_unresolved: bool,
    dry_run: bool,
    scan_parenthetical_pmids: bool,
    scan_dois: bool,
    scan_bare_dois: bool,
) -> tuple[list[dict[str, Any]], list[str]]:
    text = paragraph.text
    blocks = scan_text(
        text,
        scan_parenthetical_pmids=scan_parenthetical_pmids,
        scan_dois=scan_dois,
        scan_bare_dois=scan_bare_dois,
    )
    if not blocks:
        return [], []

    if _paragraph_has_unsafe_fields(paragraph):
        return [], [
            "Skipped PMID block in paragraph with field or hidden text content at "
            f"{location.part} paragraph {location.paragraph_index}."
        ]

    replacement_blocks = _replacement_blocks_for_text(
        text,
        scan_parenthetical_pmids=scan_parenthetical_pmids,
        scan_dois=scan_dois,
        scan_bare_dois=scan_bare_dois,
    )
    planned: list[tuple[ReplacementBlock, str]] = []
    report_replacements: list[dict[str, Any]] = []
    warnings: list[str] = []

    for block in replacement_blocks:
        replacement_text = _replacement_for_block(
            block=block,
            records_by_identifier=records_by_identifier,
            keep_pmid_text=keep_pmid_text,
            mark_unresolved=mark_unresolved,
        )
        if replacement_text is None:
            identifier_label = _replacement_block_label(block)
            warnings.append(
                f"Left unresolved {identifier_label} block unchanged at {location.part} paragraph "
                f"{location.paragraph_index}: {block.original_text}"
            )
            continue
        planned.append((block, replacement_text))
        report_replacements.append(
            {
                "original_text": block.original_text,
                "replacement_text": replacement_text,
                "pmids": list(block.pmids),
                "dois": list(block.dois),
                "identifiers": _block_identifier_report(block),
                "kind": block.kind,
                "source": block.source,
                "location": location.as_report_dict(),
            }
        )

    if not dry_run:
        for block, replacement_text in sorted(planned, key=lambda item: item[0].start, reverse=True):
            _replace_paragraph_range(paragraph, block.start, block.end, replacement_text)

    return report_replacements, warnings


def _insert_comment_pmids_at_anchors(
    *,
    document: DocxDocument,
    options: ReplacementOptions,
    records_by_identifier: dict[IdentifierKey, ReferenceRecord],
    context: ReferenceSectionContext,
) -> tuple[list[dict[str, Any]], list[str]]:
    anchor_locations = context.anchor_locations
    planned_by_comment: dict[int, list[tuple[ReplacementBlock, str, TextLocation]]] = {}
    warnings: list[str] = []

    for paragraph, comment_location in _iter_comment_paragraphs(document, options):
        if comment_location.comment_id in context.skipped_comment_ids:
            continue
        blocks = _replacement_blocks_for_text(
            paragraph.text,
            scan_parenthetical_pmids=options.scan_parenthetical_pmids,
            scan_dois=options.scan_dois,
            scan_bare_dois=options.scan_bare_dois,
        )
        if not blocks:
            continue
        if _paragraph_has_unsafe_fields(paragraph):
            warnings.append(
                "Skipped PMID block in comment with field or hidden text content at "
                f"comment {comment_location.comment_id} paragraph {comment_location.paragraph_index}."
            )
            continue
        if comment_location.comment_id is None:
            continue
        for block in blocks:
            replacement_text = _replacement_for_block(
                block=block,
                records_by_identifier=records_by_identifier,
                keep_pmid_text=options.keep_pmid_text,
                mark_unresolved=options.mark_unresolved,
            )
            if replacement_text is None:
                identifier_label = _replacement_block_label(block)
                warnings.append(
                    f"Left unresolved {identifier_label} block unchanged in comment {comment_location.comment_id}: "
                    f"{block.original_text}"
                )
                continue
            planned_by_comment.setdefault(comment_location.comment_id, []).append(
                (block, replacement_text, comment_location)
            )

    report_replacements: list[dict[str, Any]] = []
    for comment_id, planned in planned_by_comment.items():
        anchor = anchor_locations.get(comment_id)
        if anchor is None:
            identifiers = sorted(
                {identifier for block, _, _ in planned for _, identifier in block.identifier_keys}
            )
            warnings.append(
                f"Could not find document anchor for comment {comment_id}; identifiers not inserted: "
                + ", ".join(identifiers)
            )
            continue

        anchor_paragraph, anchor_location = anchor
        insertion_text = " " + " ".join(replacement_text for _, replacement_text, _ in planned)
        if not options.dry_run:
            _insert_text_after_comment_reference(anchor_paragraph, comment_id, insertion_text)

        for block, replacement_text, comment_location in planned:
            location = anchor_location.as_report_dict()
            location["comment_id"] = comment_id
            location["source_part"] = "comment"
            location["source_paragraph_index"] = comment_location.paragraph_index
            report_replacements.append(
                {
                    "original_text": block.original_text,
                    "replacement_text": replacement_text,
                    "pmids": list(block.pmids),
                    "dois": list(block.dois),
                    "identifiers": _block_identifier_report(block),
                    "kind": block.kind,
                    "source": block.source,
                    "location": location,
                }
            )

    return report_replacements, warnings


def _comment_anchor_locations(
    document: DocxDocument,
    options: ReplacementOptions,
) -> dict[int, tuple[Paragraph, TextLocation]]:
    return _reference_section_context(document, options).anchor_locations


def _comment_ids_in_paragraph(paragraph: Paragraph) -> Iterator[int]:
    for ref in paragraph._p.iter(qn("w:commentReference")):
        value = ref.get(qn("w:id"))
        if value is not None and value.isdigit():
            yield int(value)
    for ref in paragraph._p.iter(qn("w:commentRangeEnd")):
        value = ref.get(qn("w:id"))
        if value is not None and value.isdigit():
            yield int(value)


def _insert_text_after_comment_reference(
    paragraph: Paragraph,
    comment_id: int,
    text: str,
) -> None:
    anchor_element = _find_comment_anchor_element(paragraph, comment_id)
    if anchor_element is None:
        raise WordProcessingError(f"Could not find comment anchor for comment {comment_id}.")

    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    run.append(text_element)
    anchor_element.addnext(run)


def _find_comment_anchor_element(paragraph: Paragraph, comment_id: int):
    id_text = str(comment_id)
    for ref in paragraph._p.iter(qn("w:commentReference")):
        if ref.get(qn("w:id")) == id_text:
            return ref.getparent()
    for ref in paragraph._p.iter(qn("w:commentRangeEnd")):
        if ref.get(qn("w:id")) == id_text:
            return ref
    return None


def _replacement_for_block(
    *,
    block: ReplacementBlock,
    records_by_identifier: dict[IdentifierKey, ReferenceRecord],
    keep_pmid_text: bool,
    mark_unresolved: bool,
) -> str | None:
    resolved_records: list[ReferenceRecord] = []
    seen_record_keys: set[str] = set()
    unresolved_by_kind: dict[IdentifierKind, list[str]] = {"pmid": [], "doi": []}

    for identifier_key in block.identifier_keys:
        record = records_by_identifier.get(identifier_key)
        if record is None:
            unresolved_by_kind[identifier_key[0]].append(identifier_key[1])
            continue
        if record.citation_key in seen_record_keys:
            continue
        seen_record_keys.add(record.citation_key)
        resolved_records.append(record)

    replacement_parts: list[str] = []
    if resolved_records:
        replacement_parts.append(make_temporary_citation(resolved_records))

    unresolved_total = sum(len(values) for values in unresolved_by_kind.values())
    if unresolved_total:
        if mark_unresolved or resolved_records:
            for kind, values in unresolved_by_kind.items():
                if not values:
                    continue
                label = _identifier_label(kind, plural=len(values) > 1)
                replacement_parts.append(f"[unresolved {label}: {', '.join(values)}]")
        else:
            return None

    replacement_text = " ".join(replacement_parts)
    if keep_pmid_text:
        replacement_text = f"{block.original_text} {replacement_text}"
    return replacement_text


def _replacement_blocks_for_text(
    text: str,
    *,
    scan_parenthetical_pmids: bool,
    scan_dois: bool,
    scan_bare_dois: bool,
) -> list[ReplacementBlock]:
    identifier_blocks = scan_text(
        text,
        scan_parenthetical_pmids=scan_parenthetical_pmids,
        scan_dois=scan_dois,
        scan_bare_dois=scan_bare_dois,
    )
    if not identifier_blocks:
        return []

    grouped: list[ReplacementBlock] = []
    used_block_indexes: set[int] = set()
    for start, end in _identifier_only_wrappers(text):
        inside_indexes = [
            index
            for index, block in enumerate(identifier_blocks)
            if _block_inside_wrapper(block, start, end)
        ]
        if not inside_indexes or any(index in used_block_indexes for index in inside_indexes):
            continue
        inside_blocks = tuple(identifier_blocks[index] for index in inside_indexes)
        if not _wrapper_contains_only_identifiers(text, start, end, inside_blocks):
            continue
        grouped.append(
            ReplacementBlock(
                original_text=text[start:end],
                start=start,
                end=end,
                blocks=inside_blocks,
                source="wrapper",
            )
        )
        used_block_indexes.update(inside_indexes)

    for run in _adjacent_identifier_runs(text, identifier_blocks, used_block_indexes):
        if len(run) > 1:
            grouped.append(
                ReplacementBlock(
                    original_text=text[run[0].start : run[-1].end],
                    start=run[0].start,
                    end=run[-1].end,
                    blocks=tuple(run),
                    source="identifier_run",
                )
            )
            continue
        block = run[0]
        grouped.append(
            ReplacementBlock(
                original_text=block.original_text,
                start=block.start,
                end=block.end,
                blocks=(block,),
                source=block.source,
            )
        )

    return sorted(grouped, key=lambda block: block.start)


def _identifier_only_wrappers(text: str) -> Iterator[tuple[int, int]]:
    for match in re.finditer(r"\[[^\[\]]+\]", text):
        yield match.start(), match.end()
    for match in re.finditer(r"\([^()]+\)", text):
        yield match.start(), match.end()


def _adjacent_identifier_runs(
    text: str,
    blocks: list[IdentifierBlock],
    used_block_indexes: set[int],
) -> Iterator[list[IdentifierBlock]]:
    available = [
        block for index, block in enumerate(blocks)
        if index not in used_block_indexes
    ]
    if not available:
        return

    run = [available[0]]
    for block in available[1:]:
        separator = text[run[-1].end : block.start]
        if _should_group_adjacent(run[-1], block) and re.fullmatch(r"[\s,;]+", separator or ""):
            run.append(block)
            continue
        yield run
        run = [block]
    yield run


def _should_group_adjacent(left: IdentifierBlock, right: IdentifierBlock) -> bool:
    return left.source != "parenthetical" and right.source != "parenthetical"


def _block_inside_wrapper(block: IdentifierBlock, start: int, end: int) -> bool:
    return block.start >= start and block.end <= end


def _wrapper_contains_only_identifiers(
    text: str,
    start: int,
    end: int,
    blocks: tuple[IdentifierBlock, ...],
) -> bool:
    inner_start = start + 1
    inner_end = end - 1
    cursor = inner_start
    remaining: list[str] = []
    for block in sorted(blocks, key=lambda value: value.start):
        block_start = max(block.start, inner_start)
        block_end = min(block.end, inner_end)
        if block_start > cursor:
            remaining.append(text[cursor:block_start])
        cursor = max(cursor, block_end)
    if cursor < inner_end:
        remaining.append(text[cursor:inner_end])
    return re.fullmatch(r"[\s,;]*", "".join(remaining)) is not None


def _validate_replacement_citations(replacements: list[dict[str, Any]]) -> None:
    bad_replacements = [
        replacement["replacement_text"]
        for replacement in replacements
        if "#PMID-" in replacement.get("replacement_text", "")
        or "#DOI-" in replacement.get("replacement_text", "")
    ]
    if bad_replacements:
        raise WordProcessingError(
            "Generated temporary citation used the EndNote record-number marker "
            "with an app-controlled PMID/DOI key."
        )


def _replacement_block_label(block: ReplacementBlock) -> str:
    kinds = {kind for kind, _ in block.identifier_keys}
    if kinds == {"pmid"}:
        return "PMID"
    if kinds == {"doi"}:
        return "DOI"
    return "identifier"


def _identifier_label(kind: IdentifierKind, *, plural: bool = False) -> str:
    label = kind.upper()
    if plural:
        return label + "s"
    return label


def _block_identifier_report(block: ReplacementBlock) -> list[dict[str, str]]:
    report: list[dict[str, str]] = []
    for identifier_block in block.blocks:
        report.extend(
            {
                "kind": identifier_block.kind,
                "normalized": identifier,
                "source": identifier_block.source,
            }
            for identifier in identifier_block.identifiers
        )
    return report


def _skipped_identifier_reports(
    blocks: list[IdentifierBlock],
    location: TextLocation,
    *,
    reason: str,
) -> list[dict[str, Any]]:
    skipped: list[dict[str, Any]] = []
    for block in blocks:
        for identifier in block.identifiers:
            skipped.append(
                {
                    "original_text": block.original_text,
                    "kind": block.kind,
                    "normalized": identifier,
                    "source": block.source,
                    "location": location.as_report_dict(),
                    "reason": reason,
                }
            )
    return skipped


def _replace_paragraph_range(paragraph: Paragraph, start: int, end: int, replacement: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.text = paragraph.text[:start] + replacement + paragraph.text[end:]
        return

    spans: list[tuple[int, int, Any]] = []
    cursor = 0
    for run in runs:
        run_text = run.text
        run_start = cursor
        run_end = cursor + len(run_text)
        spans.append((run_start, run_end, run))
        cursor = run_end

    overlapping = [
        (run_start, run_end, run)
        for run_start, run_end, run in spans
        if run_start < end and run_end > start
    ]
    if not overlapping:
        raise WordProcessingError(
            f"Could not map paragraph replacement range {start}:{end} to Word runs."
        )

    first_start, first_end, first_run = overlapping[0]
    last_start, last_end, last_run = overlapping[-1]
    before = first_run.text[: max(0, start - first_start)]
    after = last_run.text[max(0, end - last_start) :]

    if first_run is last_run:
        first_run.text = before + replacement + after
        return

    first_run.text = before + replacement
    for _, _, run in overlapping[1:-1]:
        run.text = ""
    last_run.text = after


def _create_backup(input_docx: Path) -> Path:
    backup_path = input_docx.with_name(f"{input_docx.stem}.pmid2endnote.backup{input_docx.suffix}")
    if backup_path.exists():
        counter = 1
        while True:
            candidate = input_docx.with_name(
                f"{input_docx.stem}.pmid2endnote.backup.{counter}{input_docx.suffix}"
            )
            if not candidate.exists():
                backup_path = candidate
                break
            counter += 1
    shutil.copy2(input_docx, backup_path)
    return backup_path
