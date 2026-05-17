from pathlib import Path

from docx import Document

from pmid2endnote import app
from pmid2endnote.app import ProcessingOptions, process_document
from pmid2endnote.pubmed import PubMedRecord


def _save_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


class FakePubMedClient:
    records = {
        "6426050": PubMedRecord("6426050", "Petersen", "1983"),
        "38408451": PubMedRecord(
            "38408451",
            "Wijesundara",
            "2024",
            doi="10.1021/acs.chemrev.3c00409",
            doi_values=("10.1021/acs.chemrev.3c00409",),
        ),
        "36737783": PubMedRecord(
            "36737783",
            "Stillman",
            "2023",
            doi="10.1186/s12951-023-01782-w",
            doi_values=("10.1186/s12951-023-01782-w",),
        ),
        "30826373": PubMedRecord("30826373", "Zhong", "2019"),
        "34866536": PubMedRecord("34866536", "Li", "2021"),
    }
    doi_to_pmids = {
        "10.1021/acs.chemrev.3c00409": ["38408451"],
        "10.1186/s12951-023-01782-w": ["36737783"],
    }
    nbib_text = None

    def __init__(self, *, email: str, api_key: str | None = None) -> None:
        self.email = email
        self.api_key = api_key

    def fetch_records(self, pmids):
            return {pmid: self.records[pmid] for pmid in pmids if pmid in self.records}

    def search_pmids_by_doi(self, doi):
        return self.doi_to_pmids.get(doi, [])

    def fetch_nbib(self, pmids):
        if self.nbib_text is not None:
            return self.nbib_text
        return "\n\n".join(f"PMID- {pmid}\nTI  - Example" for pmid in pmids)


def test_non_dry_run_writes_valid_enw_before_docx_and_reports_status(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setattr(app, "PubMedClient", FakePubMedClient)
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "input.endnote.docx"
    nbib_file = tmp_path / "input.references.nbib"
    enw_file = tmp_path / "input.endnote-import.enw"
    report_file = tmp_path / "input.report.json"
    _save_docx(input_docx, "See (6426050).")

    result = process_document(
        ProcessingOptions(
            input_docx=input_docx,
            email="test@example.edu",
            output_docx=output_docx,
            nbib_file=nbib_file,
            enw_file=enw_file,
            report_file=report_file,
            scan_parenthetical_pmids=True,
            save_email=False,
        )
    )

    assert result.exit_code == 0
    assert enw_file.exists()
    assert nbib_file.exists()
    assert output_docx.exists()
    assert "%M PMID-6426050" in enw_file.read_text(encoding="utf-8")
    assert result.report["create_backup"] is False
    assert not any("Created backup copy" in warning for warning in result.report["warnings"])
    assert result.report["pmid_statuses"] == [
        {
            "pmid": "6426050",
            "resolved": True,
            "included_in_nbib": True,
            "replacement_count": 1,
            "sources": ["parenthetical"],
            "warnings": [],
        }
    ]
    assert result.report["identifier_statuses"][0]["citation_key"] == "PMID-6426050"
    assert str(enw_file) in result.messages[-1]
    assert "EndNote Import option" in result.messages[-1]
    assert "Accession Number" not in result.messages[-1]


def test_empty_auxiliary_nbib_does_not_block_docx_when_enw_is_valid(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(app, "PubMedClient", FakePubMedClient)
    monkeypatch.setattr(FakePubMedClient, "nbib_text", "")
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "input.endnote.docx"
    nbib_file = tmp_path / "input.references.nbib"
    report_file = tmp_path / "input.report.json"
    _save_docx(input_docx, "PMID: 6426050")

    result = process_document(
        ProcessingOptions(
            input_docx=input_docx,
            email="test@example.edu",
            output_docx=output_docx,
            nbib_file=nbib_file,
            report_file=report_file,
            save_email=False,
        )
    )

    assert result.exit_code == 0
    assert nbib_file.exists()
    assert output_docx.exists()
    assert any("empty" in warning for warning in result.report["warnings"])


def test_enw_validation_failure_blocks_docx(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(app, "PubMedClient", FakePubMedClient)
    monkeypatch.setattr(
        app,
        "validate_enw_file",
        lambda _path, _records: (set(), ["EndNote Tagged Import file was not written: missing"]),
    )
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "input.endnote.docx"
    _save_docx(input_docx, "PMID: 6426050")

    result = process_document(
        ProcessingOptions(
            input_docx=input_docx,
            email="test@example.edu",
            output_docx=output_docx,
            nbib_file=tmp_path / "input.references.nbib",
            report_file=tmp_path / "input.report.json",
            save_email=False,
        )
    )

    assert result.exit_code == 2
    assert not output_docx.exists()
    assert "not written" in result.report["errors"][0]


def test_dry_run_does_not_require_written_nbib_or_docx(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(app, "PubMedClient", FakePubMedClient)
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "input.endnote.docx"
    nbib_file = tmp_path / "input.references.nbib"
    _save_docx(input_docx, "PMID: 6426050")

    result = process_document(
        ProcessingOptions(
            input_docx=input_docx,
            email="test@example.edu",
            output_docx=output_docx,
            nbib_file=nbib_file,
            report_file=tmp_path / "input.report.json",
            dry_run=True,
            save_email=False,
        )
    )

    assert result.exit_code == 0
    assert not nbib_file.exists()
    assert not output_docx.exists()
    assert not result.enw_file.exists()
    assert result.report["pmid_statuses"][0]["replacement_count"] == 1


def test_reference_section_skipped_identifiers_do_not_enter_import_or_statuses(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setattr(app, "PubMedClient", FakePubMedClient)
    input_docx = tmp_path / "stress.docx"
    output_docx = tmp_path / "stress.endnote.docx"
    nbib_file = tmp_path / "stress.references.nbib"
    enw_file = tmp_path / "stress.endnote-import.enw"
    report_file = tmp_path / "stress.report.json"
    document = Document()
    document.add_paragraph("Body PMID: 38408451 and DOI: 10.1186/s12951-023-01782-w.")
    document.add_paragraph("Parenthetical stress (30826373; 34866536).")
    document.add_paragraph("References")
    document.add_paragraph("1. Skipped repeat. PMID: 38408451. DOI: 10.1021/acs.chemrev.3c00409.")
    document.add_paragraph("2. Skipped DOI-only-looking item. DOI: 10.1002/adfm.201600650.")
    document.save(input_docx)

    result = process_document(
        ProcessingOptions(
            input_docx=input_docx,
            email="test@example.edu",
            output_docx=output_docx,
            nbib_file=nbib_file,
            enw_file=enw_file,
            report_file=report_file,
            scan_parenthetical_pmids=True,
            save_email=False,
        )
    )

    assert result.exit_code == 0
    assert {item["normalized"] for item in result.report["skipped_identifiers"]} == {
        "38408451",
        "10.1021/acs.chemrev.3c00409",
        "10.1002/adfm.201600650",
    }
    assert result.report["reference_section_start"] == {"part": "body", "paragraph_index": 2}
    assert {item["normalized"] for item in result.report["unique_identifiers"]} == {
        "38408451",
        "10.1186/s12951-023-01782-w",
        "30826373",
        "34866536",
    }
    assert {item["normalized"] for item in result.report["identifier_statuses"]} == {
        "38408451",
        "10.1186/s12951-023-01782-w",
        "30826373",
        "34866536",
    }
    enw_text = enw_file.read_text(encoding="utf-8")
    assert "PMID-38408451" in enw_text
    assert "PMID-36737783" in enw_text
    assert "PMID-30826373" in enw_text
    assert "PMID-34866536" in enw_text
    assert "10.1002/adfm.201600650" not in enw_text
    output_document = Document(output_docx)
    assert output_document.paragraphs[3].text == (
        "1. Skipped repeat. PMID: 38408451. DOI: 10.1021/acs.chemrev.3c00409."
    )
