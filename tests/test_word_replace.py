from pathlib import Path

from docx import Document

from pmid2endnote.pubmed import PubMedRecord
from pmid2endnote.word import ReplacementOptions, replace_pmids_in_docx, scan_docx


def _save_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def _read_paragraph_text(path: Path, index: int = 0) -> str:
    return Document(path).paragraphs[index].text


def test_word_replacement_preserves_non_pmid_text(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["Before PMID: 12345678 after."])

    result = replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"12345678": PubMedRecord("12345678", "Smith", "2024")},
    )

    assert _read_paragraph_text(output_docx) == "Before {Smith, 2024 #PMID-12345678} after."
    assert result.replacements[0]["original_text"] == "PMID: 12345678"
    assert result.backup_path is not None
    assert result.backup_path.exists()


def test_word_replacement_preserves_first_replaced_run_formatting(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("See ")
    pmid_run = paragraph.add_run("PMID: 12345678")
    pmid_run.bold = True
    paragraph.add_run(" now.")
    document.save(input_docx)

    replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"12345678": PubMedRecord("12345678", "Smith", "2024")},
    )

    output_paragraph = Document(output_docx).paragraphs[0]
    assert output_paragraph.text == "See {Smith, 2024 #PMID-12345678} now."
    assert output_paragraph.runs[1].bold is True


def test_unresolved_pmid_handling_leaves_unresolved_text(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["PMIDs: 12345678, 99999999"])

    replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"12345678": PubMedRecord("12345678", "Smith", "2024")},
    )

    assert _read_paragraph_text(output_docx) == "{Smith, 2024 #PMID-12345678} [unresolved PMID: 99999999]"


def test_dry_run_reports_without_writing_output(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["PMID: 12345678"])

    result = replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"12345678": PubMedRecord("12345678", "Smith", "2024")},
        options=ReplacementOptions(dry_run=True),
    )

    assert result.replacements[0]["replacement_text"] == "{Smith, 2024 #PMID-12345678}"
    assert not output_docx.exists()
    assert _read_paragraph_text(input_docx) == "PMID: 12345678"


def test_scan_docx_finds_tables_by_default(tmp_path: Path) -> None:
    input_docx = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "PMID: 12345678"
    document.save(input_docx)

    scan = scan_docx(input_docx)

    assert scan.unique_pmids == ["12345678"]
    assert scan.occurrences[0].location.part == "table"


def test_scan_docx_finds_pmids_in_comments_by_default(tmp_path: Path) -> None:
    input_docx = tmp_path / "comment.docx"
    document = Document()
    paragraph = document.add_paragraph("Anchor text")
    document.add_comment(paragraph.runs[0], text="Please cite PMID: 12345678", author="Reviewer")
    document.save(input_docx)

    scan = scan_docx(input_docx)

    assert scan.unique_pmids == ["12345678"]
    assert scan.occurrences[0].location.part == "comment"
    assert scan.occurrences[0].location.comment_id == 0


def test_scan_docx_can_skip_comments(tmp_path: Path) -> None:
    input_docx = tmp_path / "comment.docx"
    document = Document()
    paragraph = document.add_paragraph("Anchor text")
    document.add_comment(paragraph.runs[0], text="Please cite PMID: 12345678", author="Reviewer")
    document.save(input_docx)

    scan = scan_docx(input_docx, ReplacementOptions(include_comments=False))

    assert scan.unique_pmids == []


def test_parenthetical_single_pmid_replacement(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["Before (6426050) after."])

    replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"6426050": PubMedRecord("6426050", "Petersen", "1983")},
        options=ReplacementOptions(scan_parenthetical_pmids=True),
    )

    assert _read_paragraph_text(output_docx) == "Before {Petersen, 1983 #PMID-6426050} after."


def test_word_replacement_inserts_comment_pmids_at_anchor(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("before ")
    anchor = paragraph.add_run("stimulation")
    paragraph.add_run(" after")
    document.add_comment(anchor, text="Please cite PMID: 12345678", author="Reviewer")
    document.save(input_docx)

    result = replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"12345678": PubMedRecord("12345678", "Smith", "2024")},
    )

    output_document = Document(output_docx)
    assert output_document.paragraphs[0].text == "before stimulation {Smith, 2024 #PMID-12345678} after"
    assert output_document.comments.get(0).text == "Please cite PMID: 12345678"
    assert result.replacements[0]["location"]["part"] == "body"
    assert result.replacements[0]["location"]["comment_id"] == 0
    assert result.replacements[0]["location"]["source_part"] == "comment"


def test_word_replacement_can_skip_comment_anchor_insertions(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    document = Document()
    paragraph = document.add_paragraph("Anchor text")
    document.add_comment(paragraph.runs[0], text="Please cite PMID: 12345678", author="Reviewer")
    document.save(input_docx)

    result = replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"12345678": PubMedRecord("12345678", "Smith", "2024")},
        options=ReplacementOptions(include_comments=False),
    )

    assert Document(output_docx).paragraphs[0].text == "Anchor text"
    assert result.replacements == []


def test_parenthetical_multi_pmid_replacement(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["Before (6426050, 104929) after."])

    replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={
            "6426050": PubMedRecord("6426050", "Petersen", "1983"),
            "104929": PubMedRecord("104929", "Petersen", "1978"),
        },
        options=ReplacementOptions(scan_parenthetical_pmids=True),
    )

    assert (
        _read_paragraph_text(output_docx)
        == "Before {Petersen, 1983 #PMID-6426050;Petersen, 1978 #PMID-104929} after."
    )


def test_parenthetical_unresolved_pmid_stays_outside_braces(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["Before (6426050, 999999999) after."])

    replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={"6426050": PubMedRecord("6426050", "Petersen", "1983")},
        options=ReplacementOptions(scan_parenthetical_pmids=True),
    )

    assert (
        _read_paragraph_text(output_docx)
        == "Before {Petersen, 1983 #PMID-6426050} [unresolved PMID: 999999999] after."
    )


def test_parenthetical_mode_preserves_surrounding_punctuation(tmp_path: Path) -> None:
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    _save_docx(input_docx, ["activity. (6426050), (104929), (9036716)."])

    replace_pmids_in_docx(
        input_docx=input_docx,
        output_docx=output_docx,
        records_by_pmid={
            "6426050": PubMedRecord("6426050", "A", "1983"),
            "104929": PubMedRecord("104929", "B", "1978"),
            "9036716": PubMedRecord("9036716", "C", "1997"),
        },
        options=ReplacementOptions(scan_parenthetical_pmids=True),
    )

    assert (
        _read_paragraph_text(output_docx)
        == "activity. {A, 1983 #PMID-6426050}, {B, 1978 #PMID-104929}, {C, 1997 #PMID-9036716}."
    )
